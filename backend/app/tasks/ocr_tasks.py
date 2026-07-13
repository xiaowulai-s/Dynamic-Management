import os
import json
import fitz  # PyMuPDF
from docx import Document
import easyocr
import numpy as np
from PIL import Image
from celery import shared_task
from app.config import settings
from app.utils.file_utils import get_file_path


@shared_task(bind=True, name="ocr_tasks.process_uploaded_files")
def process_uploaded_files(self, file_paths: list) -> dict:
    """
    处理上传的文件，进行OCR识别

    Args:
        file_paths: 文件路径列表

    Returns:
        dict: {
            "task_id": str,
            "status": "completed" | "failed",
            "results": [
                {
                    "file_path": str,
                    "success": bool,
                    "text": str,
                    "structured_data": dict
                }
            ]
        }
    """
    results = []
    total_files = len(file_paths)

    for index, file_path in enumerate(file_paths):
        try:
            # 更新进度
            self.update_state(
                state="PROGRESS",
                meta={
                    "current": index + 1,
                    "total": total_files,
                    "status": f"正在处理文件 {index + 1}/{total_files}"
                }
            )

            # 识别文件
            file_result = recognize_file(file_path)
            results.append(file_result)

        except Exception as e:
            results.append({
                "file_path": file_path,
                "success": False,
                "error": str(e),
                "text": "",
                "structured_data": {}
            })

    return {
        "task_id": self.request.id,
        "status": "completed",
        "results": results
    }


def recognize_file(file_path: str) -> dict:
    """
    识别单个文件

    Args:
        file_path: 文件路径

    Returns:
        dict: 识别结果
    """
    full_path = get_file_path(file_path)
    file_ext = os.path.splitext(file_path)[1].lower()

    # 根据文件类型选择识别方式
    if file_ext == ".pdf":
        return recognize_pdf(full_path)
    elif file_ext == ".docx":
        return recognize_docx(full_path)
    elif file_ext in [".jpg", ".jpeg", ".png"]:
        return recognize_image(full_path)
    else:
        return {
            "file_path": file_path,
            "success": False,
            "error": f"不支持的文件格式：{file_ext}",
            "text": "",
            "structured_data": {}
        }


def recognize_pdf(pdf_path: str) -> dict:
    """
    识别PDF文件

    Args:
        pdf_path: PDF文件路径

    Returns:
        dict: 识别结果
    """
    try:
        # 方法1：尝试直接提取文本（如果PDF是文字型）
        doc = fitz.open(pdf_path)
        text = ""
        structured_data = {}

        for page_num in range(len(doc)):
            page = doc[page_num]
            text += page.get_text()

        # 如果提取的文本太少，可能是扫描件，使用OCR
        if len(text.strip()) < 50:
            # 将PDF转换为图片
            for page_num in range(min(len(doc), 3)):  # 只处理前3页
                page = doc[page_num]
                mat = fitz.Matrix(2, 2)  # 2倍缩放以提高OCR准确率
                pix = page.get_pixmap(matrix=mat)
                img_path = f"temp_{page_num}.png"
                pix.save(img_path)

                # OCR识别
                reader = easyocr.Reader(['ch_sim', 'en'])
                result = reader.readtext(img_path)

                page_text = ""
                for detection in result:
                    page_text += detection[1] + "\n"

                text += page_text + "\n"

                # 删除临时文件
                if os.path.exists(img_path):
                    os.remove(img_path)

        doc.close()

        # 解析结构化数据
        structured_data = parse_log_text(text)

        return {
            "file_path": pdf_path,
            "success": True,
            "text": text,
            "structured_data": structured_data
        }

    except Exception as e:
        return {
            "file_path": pdf_path,
            "success": False,
            "error": str(e),
            "text": "",
            "structured_data": {}
        }


def recognize_docx(docx_path: str) -> dict:
    """
    识别Word文档

    Args:
        docx_path: Word文档路径

    Returns:
        dict: 识别结果
    """
    try:
        doc = Document(docx_path)
        text = ""

        # 提取段落文本
        for para in doc.paragraphs:
            text += para.text + "\n"

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"

        # 解析结构化数据
        structured_data = parse_log_text(text)

        return {
            "file_path": docx_path,
            "success": True,
            "text": text,
            "structured_data": structured_data
        }

    except Exception as e:
        return {
            "file_path": docx_path,
            "success": False,
            "error": str(e),
            "text": "",
            "structured_data": {}
        }


def recognize_image(image_path: str) -> dict:
    """
    识别图片

    Args:
        image_path: 图片路径

    Returns:
        dict: 识别结果
    """
    try:
        # 使用PaddleOCR识别
        reader = easyocr.Reader(['ch_sim', 'en'])
        result = reader.readtext(image_path)

        text = ""
        for detection in result:
            text += detection[1] + "\n"

        # 解析结构化数据
        structured_data = parse_log_text(text)

        return {
            "file_path": image_path,
            "success": True,
            "text": text,
            "structured_data": structured_data
        }

    except Exception as e:
        return {
            "file_path": image_path,
            "success": False,
            "error": str(e),
            "text": "",
            "structured_data": {}
        }


def parse_log_text(text: str) -> dict:
    """
    从识别的文本中解析日志结构化数据

    Args:
        text: 识别的文本

    Returns:
        dict: 结构化数据
    """
    import re

    data = {}

    # 提取设备编号（常见格式：设备编号：XXX 或 编号：XXX 或 NO：XXX）
    patterns = {
        "equipment_code": r"设备编号[：:]\s*(\S+)",
        "equipment_name": r"设备名称[：:]\s*(\S+)",
        "date": r"日期[：:]\s*(\d{4}[-年/]\d{1,2}[-月/]\d{1,2}[日]?)",
        "operator": r"操作人[：:]\s*(\S+)",
        "description": r"描述[：:]\s*(.+)"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            data[key] = match.group(1).strip()

    # 尝试解析日期
    if "date" in data:
        date_str = data["date"]
        # 统一转换为 YYYY-MM-DD 格式
        date_str = date_str.replace("年", "-").replace("月", "-").replace("日", "").replace("/", "-")
        data["date"] = date_str

    # 检测日志类型
    if "安装" in text or "验收" in text:
        data["log_type"] = "installation"
    elif "维修" in text or "修理" in text:
        data["log_type"] = "repair"
    elif "报废" in text:
        data["log_type"] = "scrap"
    elif "巡检" in text:
        data["log_type"] = "inspection"
    elif "保养" in text:
        data["log_type"] = "maintenance"
    elif "报修" in text or "故障" in text:
        data["log_type"] = "fault"
    elif "配件" in text or "更换" in text:
        data["log_type"] = "parts"
    elif "校准" in text:
        data["log_type"] = "calibration"
    else:
        data["log_type"] = "unknown"

    return data
